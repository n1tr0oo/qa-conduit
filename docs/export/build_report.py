"""
Generates a formally formatted Word document for Assignment 3 report.
Run: python docs/export/build_report.py
Output: docs/export/assignment3-report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BLACK = RGBColor(0, 0, 0)
HEADER_BG = "2E74B5"   # dark blue for table headers
ALT_BG    = "DEEAF1"   # light blue alternate rows

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=BLACK):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color

def para_font(para, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    para.alignment = align
    for run in para.runs:
        set_font(run, size, bold)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name  = "Times New Roman"
        run.font.color.rgb = BLACK
        run.font.bold  = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p

def add_para(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph(text)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        set_font(run, size, bold)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.font.name  = "Times New Roman"
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = ALT_BG if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(str(val))
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(9)
            run.font.bold  = str(val).startswith("**") or str(val).lower() in ("overall", "result", "✓ pass", "✗ fail")
            run.font.color.rgb = BLACK

    # column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── Title page ────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ASSIGNMENT 3: EXPERIMENTAL ENGINEERING")
r.font.name = "Times New Roman"; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = BLACK

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Performance / Mutation / Chaos Testing")
r2.font.name = "Times New Roman"; r2.font.size = Pt(14); r2.font.bold = False; r2.font.color.rgb = BLACK

doc.add_paragraph()
for line in [
    "Project: Conduit RealWorld App",
    "Repository: https://github.com/n1tr0oo/qa-conduit",
    "Date: 2026-04-24",
    "Stack: React 19 / Express.js 5 / Sequelize 6 / PostgreSQL 15 / Node.js 20",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = "Times New Roman"; r.font.size = Pt(11); r.font.color.rgb = BLACK

doc.add_page_break()

# ── 1. System Description ─────────────────────────────────────────────────────
add_heading(doc, "1. System Description and High-Risk Modules", 1)
add_para(doc,
    "Conduit is a full-stack blogging platform implementing the RealWorld API specification. "
    "The system provides user authentication via JWT tokens, article management with slug-based routing, "
    "a global article feed with tag filtering and pagination, and social features (comments, favorites, follows).")
add_para(doc,
    "Three high-risk modules were selected from the midterm Risk-Based Test Matrix (RBTM) as targets "
    "for all experimental testing:")

add_table(doc,
    ["#", "Module", "Risk Level", "Justification"],
    [
        ["1", "User Authentication", "HIGH", "JWT tokens gate all protected endpoints; failure blocks entire authenticated user flow"],
        ["2", "Article Management", "HIGH", "Core business function; slug generation affects routing; CRUD failures impact all content"],
        ["3", "Global Feed & Pagination", "HIGH", "Primary entry point for all users; broken feed renders the application unusable"],
    ],
    col_widths=[1.0, 4.0, 2.5, 8.5]
)

# ── 2. Performance Testing ────────────────────────────────────────────────────
add_heading(doc, "2. Performance Testing", 1)

add_heading(doc, "2.1 Test Plan", 2)
add_para(doc, "Tool: Newman 6.2.2 (Postman CLI — programmatic API with parallel virtual user simulation)")
add_para(doc, "Endpoints under test (9 total across 3 high-risk modules):")

add_table(doc,
    ["ID", "Endpoint", "Module"],
    [
        ["AUTH-P01", "POST /api/users/login", "Authentication"],
        ["AUTH-P02", "GET /api/user", "Authentication"],
        ["FEED-P01", "GET /api/articles", "Feed"],
        ["FEED-P02", "GET /api/articles?limit=10&offset=0", "Feed"],
        ["FEED-P03", "GET /api/articles?tag=qa", "Feed"],
        ["FEED-P04", "GET /api/tags", "Feed"],
        ["ART-P01", "POST /api/articles", "Articles"],
        ["ART-P02", "GET /api/articles/:slug", "Articles"],
        ["ART-P03", "DELETE /api/articles/:slug", "Articles"],
    ],
    col_widths=[2.5, 7.0, 4.0]
)

add_para(doc, "Scenario definitions:")
add_table(doc,
    ["Scenario", "Virtual Users", "Iterations", "Delay", "Total Requests", "p95 Threshold"],
    [
        ["Normal load", "1", "30", "200ms", "270", "< 500ms"],
        ["Peak load", "5 (parallel)", "20 each", "100ms", "900", "< 800ms"],
        ["Spike load", "15 (parallel)", "5 each", "0ms", "675", "< 1500ms"],
        ["Endurance", "1", "100", "100ms", "900", "< 600ms"],
    ],
    col_widths=[3.5, 3.0, 2.5, 2.0, 3.5, 3.5]
)

add_heading(doc, "2.2 Performance Results", 2)

for title, rows, throughput, walltime, result in [
    ("Normal Load — 1 VU × 30 iterations — PASS ✓",
     [
        ["AUTH-P01: Login", "30", "61ms", "60ms", "65ms", "87ms", "0.0%"],
        ["AUTH-P02: Get current user", "30", "2ms", "2ms", "3ms", "5ms", "0.0%"],
        ["FEED-P01: Global feed", "30", "6ms", "6ms", "11ms", "11ms", "0.0%"],
        ["FEED-P02: Paginated feed", "30", "6ms", "6ms", "11ms", "11ms", "0.0%"],
        ["FEED-P03: Tag filter", "30", "3ms", "3ms", "5ms", "6ms", "0.0%"],
        ["FEED-P04: Tags list", "30", "2ms", "2ms", "3ms", "4ms", "0.0%"],
        ["ART-P01: Create article", "30", "6ms", "6ms", "10ms", "11ms", "0.0%"],
        ["ART-P02: Get by slug", "30", "5ms", "5ms", "7ms", "9ms", "0.0%"],
        ["ART-P03: Delete article", "30", "4ms", "4ms", "5ms", "6ms", "0.0%"],
        ["OVERALL", "270", "11ms", "5ms", "60ms", "63ms", "0.0%"],
     ], "3.2 req/s", "84.8s", "PASS"),
    ("Peak Load — 5 VU × 20 iterations — PASS ✓",
     [
        ["AUTH-P01: Login", "100", "64ms", "61ms", "95ms", "152ms", "0.0%"],
        ["AUTH-P02: Get current user", "100", "3ms", "3ms", "6ms", "12ms", "0.0%"],
        ["FEED-P01: Global feed", "100", "10ms", "9ms", "19ms", "27ms", "1.0%"],
        ["FEED-P02: Paginated feed", "100", "11ms", "10ms", "22ms", "27ms", "0.0%"],
        ["FEED-P03: Tag filter", "100", "5ms", "4ms", "8ms", "11ms", "0.0%"],
        ["FEED-P04: Tags list", "100", "3ms", "3ms", "8ms", "10ms", "0.0%"],
        ["ART-P01: Create article", "100", "10ms", "9ms", "18ms", "22ms", "0.0%"],
        ["ART-P02: Get by slug", "100", "8ms", "7ms", "13ms", "17ms", "0.0%"],
        ["ART-P03: Delete article", "100", "5ms", "5ms", "9ms", "12ms", "0.0%"],
        ["OVERALL", "900", "13ms", "7ms", "61ms", "71ms", "0.2%"],
     ], "25.8 req/s", "34.9s", "PASS"),
    ("Spike Load — 15 VU × 5 iterations — PASS ✓",
     [
        ["AUTH-P01: Login", "75", "139ms", "135ms", "267ms", "344ms", "0.0%"],
        ["AUTH-P02: Get current user", "75", "29ms", "17ms", "73ms", "196ms", "0.0%"],
        ["FEED-P01: Global feed", "75", "52ms", "43ms", "162ms", "226ms", "0.0%"],
        ["FEED-P02: Paginated feed", "75", "55ms", "47ms", "91ms", "211ms", "0.0%"],
        ["FEED-P03: Tag filter", "75", "21ms", "19ms", "40ms", "62ms", "0.0%"],
        ["FEED-P04: Tags list", "75", "12ms", "10ms", "31ms", "40ms", "0.0%"],
        ["ART-P01: Create article", "75", "46ms", "47ms", "69ms", "71ms", "0.0%"],
        ["ART-P02: Get by slug", "75", "36ms", "34ms", "59ms", "61ms", "0.0%"],
        ["ART-P03: Delete article", "75", "27ms", "24ms", "84ms", "87ms", "0.0%"],
        ["OVERALL", "675", "46ms", "32ms", "146ms", "244ms", "0.0%"],
     ], "120.5 req/s", "5.6s", "PASS"),
]:
    add_para(doc, title, bold=True, size=11)
    add_table(doc,
        ["Request", "n", "avg", "p50", "p95", "p99", "err%"],
        rows,
        col_widths=[5.5, 1.2, 1.5, 1.5, 1.5, 1.5, 1.3]
    )
    add_para(doc, f"Throughput: {throughput}  |  Wall time: {walltime}", size=10)

add_para(doc, "Endurance — 1 VU × 100 iterations — PASS ✓", bold=True)
add_table(doc,
    ["Metric", "Value"],
    [
        ["Total requests", "900"],
        ["avg response time", "11ms"],
        ["p50", "5ms"],
        ["p95", "60ms"],
        ["p99", "65ms"],
        ["Error rate", "0.0%"],
        ["Throughput", "4.5 req/s"],
        ["Wall time", "198.5s"],
        ["Latency drift", "None — no memory leak detected"],
    ],
    col_widths=[6.0, 6.0]
)

add_heading(doc, "2.3 Performance Summary", 2)
add_table(doc,
    ["Scenario", "VUs", "Requests", "avg", "p95", "Throughput", "Errors", "Result"],
    [
        ["Normal",   "1",  "270", "11ms", "60ms",  "3.2 req/s",   "0.0%", "✓ PASS"],
        ["Peak",     "5",  "900", "13ms", "61ms",  "25.8 req/s",  "0.2%", "✓ PASS"],
        ["Spike",    "15", "675", "46ms", "146ms", "120.5 req/s", "0.0%", "✓ PASS"],
        ["Endurance","1",  "900", "11ms", "60ms",  "4.5 req/s",   "0.0%", "✓ PASS"],
    ],
    col_widths=[2.8, 1.3, 2.5, 1.5, 1.8, 3.0, 1.8, 1.8]
)
add_para(doc, "p95 degradation: 60ms (normal, 1 VU) → 146ms (spike, 15 VU) — 2.4× increase under 15× load.")

add_heading(doc, "2.4 Bottleneck Analysis", 2)
add_para(doc,
    "The Authentication endpoint (POST /api/users/login) is the primary bottleneck. bcrypt password "
    "hashing with cost factor 10 introduces ~60ms CPU delay per request, capping login throughput at "
    "approximately 17 req/s regardless of load level. This is an intentional security feature, not a defect.")
add_para(doc, "Optimization recommendations:")
add_table(doc,
    ["Priority", "Module", "Issue", "Recommendation"],
    [
        ["HIGH",   "Authentication", "bcrypt throughput cap (~17 RPS)", "Add rate limiting; JWT refresh token caching"],
        ["MEDIUM", "Feed",           "N+1 Sequelize query on article list", "Eager loading; in-process LRU cache for GET /api/articles"],
        ["LOW",    "All GET",        "No HTTP caching headers", "Add Cache-Control: public, max-age=30"],
    ],
    col_widths=[2.0, 3.5, 5.0, 6.0]
)

# ── 3. Mutation Testing ───────────────────────────────────────────────────────
add_heading(doc, "3. Mutation Testing", 1)

add_heading(doc, "3.1 Mutation Plan", 2)
add_para(doc, "Tool: Stryker.js v9.6.1 (@stryker-mutator/core + @stryker-mutator/vitest-runner)")
add_para(doc, "Target: 5 pure backend helper functions extracted from backend/helper/helpers.js")
add_para(doc, "Test suite: 44 vitest tests | Duration: 8 seconds")
add_para(doc, "Mutation types applied:")
add_table(doc,
    ["Mutation Type", "Description", "Example"],
    [
        ["ArithmeticOperator", "Change arithmetic operators", "+ → -"],
        ["EqualityOperator",   "Change comparison operators", ">= → >, === → !=="],
        ["LogicalOperator",    "Invert boolean logic",        "&& → ||"],
        ["MethodExpression",   "Remove method calls",         ".trim() → removed"],
        ["Regex",              "Weaken regular expressions",  "remove ^ or $ anchors"],
        ["StringLiteral",      "Change string values",        "'-' → ''"],
        ["ConditionalExpression", "Flip ternary conditions",  "x?a:b → x?b:a"],
    ],
    col_widths=[4.5, 6.5, 5.5]
)

add_heading(doc, "3.2 Mutation Score", 2)
add_table(doc,
    ["Module / Function", "Mutants Created", "Mutants Killed", "Mutants Survived", "Score (%)"],
    [
        ["slugify",                "~15", "15", "0", "100.0%"],
        ["appendTagList",          "~7",  "6",  "1", "85.7%"],
        ["validateRequiredFields", "~9",  "9",  "0", "100.0%"],
        ["isValidEmail",           "~15", "12", "3", "80.0%"],
        ["paginationParams",       "~7",  "6",  "1", "85.7%"],
        ["OVERALL",                "53",  "48", "5", "90.57%"],
    ],
    col_widths=[5.0, 3.5, 3.5, 3.5, 3.0]
)
add_para(doc, "Formula: Mutation Score = (48 / 53) × 100 = 90.57%")
add_para(doc, "Quality thresholds: break=50%, low=60%, high=80% → Score exceeds HIGH threshold ✓")

add_heading(doc, "3.3 Surviving Mutants Analysis", 2)
add_table(doc,
    ["#", "Type", "Original", "Mutant", "Root Cause"],
    [
        ["M-S1", "LogicalOperator",    "dataValues || {}",          "dataValues && {}",         "Tests always pass pre-initialized objects"],
        ["M-S2", "Regex",              "/^[^\\s@]+.../",            "removed ^ anchor",          "No test with leading-whitespace email"],
        ["M-S3", "Regex",              ".../[^\\s@]+$/",            "removed $ anchor",          "No test with trailing-garbage email"],
        ["M-S4", "MethodExpression",   "email.trim()",              "email (no trim)",           "All test emails have no surrounding whitespace"],
        ["M-S5", "EqualityOperator",   "offset >= 0",               "offset > 0",               "offset=0 returns 0 either way — equivalent mutant"],
    ],
    col_widths=[1.2, 3.8, 3.5, 3.5, 5.5]
)
add_para(doc, "3 of 5 survivors (M-S2, M-S3, M-S4) are genuine test gaps. M-S1 and M-S5 are equivalent mutants.")
add_para(doc, "Recommended additions: test isValidEmail(' valid@email.com') and isValidEmail('valid@email.com ') to kill regex survivors.")

# ── 4. Chaos Testing ─────────────────────────────────────────────────────────
add_heading(doc, "4. Chaos / Fault Injection Testing", 1)

add_heading(doc, "4.1 Chaos Testing Plan", 2)
add_table(doc,
    ["Scenario", "Fault Type", "Affected Modules", "Duration"],
    [
        ["A: API Downtime",         "Connection refused (dead port :19999)", "Auth, Feed, Articles", "Instant"],
        ["B: Database Failure",     "Mock 500/503 + 8s slow write",          "Auth, Feed, Articles", "8s (slow write)"],
        ["C: Network Latency",      "Proxy injects 250–3000ms delay",        "Auth, Feed, Articles", "Per-request"],
        ["D: Resource Exhaustion",  "1–100 concurrent request flood",        "Auth, Feed, Articles", "~2 min total"],
    ],
    col_widths=[4.0, 5.5, 4.5, 2.5]
)

add_heading(doc, "4.2 Scenario A — API Downtime", 2)
add_table(doc,
    ["Scenario", "Module", "HTTP Status", "Response Time", "Graceful"],
    [
        ["AUTH-DOWN-01: Login unavailable",    "Auth",     "ECONNREFUSED", "2ms", "✓"],
        ["AUTH-DOWN-02: Register unavailable", "Auth",     "ECONNREFUSED", "1ms", "✓"],
        ["FEED-DOWN-01: Global feed",          "Feed",     "ECONNREFUSED", "0ms", "✓"],
        ["FEED-DOWN-02: Tags unavailable",     "Feed",     "ECONNREFUSED", "1ms", "✓"],
        ["ART-DOWN-01: Articles unavailable",  "Articles", "ECONNREFUSED", "1ms", "✓"],
    ],
    col_widths=[5.5, 2.5, 3.5, 3.5, 1.5]
)
add_para(doc, "Result: 5/5 PASS ✓  |  Avg response: 1ms  |  Max: 2ms  |  Availability: 0% (expected)  |  MTTR: N/A")

add_heading(doc, "4.3 Scenario B — Database Failure", 2)
add_table(doc,
    ["Scenario", "HTTP", "Response Time", "JSON Error Body", "Result"],
    [
        ["DB-FAIL-01: Login — DB down",    "500", "11ms",   "✓", "PASS ✓"],
        ["DB-FAIL-02: Register — DB down", "500", "1ms",    "✓", "PASS ✓"],
        ["DB-FAIL-03: Feed — DB down",     "503", "1ms",    "✓", "PASS ✓"],
        ["DB-FAIL-04: Tags — DB down",     "503", "1ms",    "✓", "PASS ✓"],
        ["DB-SLOW-01: Write — 8s lock",    "401", "8021ms", "✓", "PASS ✓"],
    ],
    col_widths=[5.5, 1.5, 3.5, 3.5, 2.5]
)
add_para(doc, "Result: 5/5 PASS ✓  |  MTTR (slow write): 8021ms  |  All responses use consistent {errors:{...}} JSON schema.")

add_heading(doc, "4.4 Scenario C — Network Latency Injection", 2)
add_table(doc,
    ["Scenario", "Injected Delay", "avg", "p50", "p95", "Errors", "Result"],
    [
        ["Normal baseline",    "0ms",    "35ms",   "16ms",   "76ms",   "0/9", "✓"],
        ["Light latency",      "250ms",  "252ms",  "266ms",  "320ms",  "1/9", "✓"],
        ["Moderate latency",   "1000ms", "917ms",  "1010ms", "1074ms", "1/9", "✓"],
        ["Severe latency",     "3000ms", "2699ms", "3018ms", "3076ms", "1/9", "✓"],
    ],
    col_widths=[3.5, 3.0, 2.2, 2.2, 2.2, 2.0, 1.4]
)
add_para(doc, "Availability under severe latency: ~89%  |  p95 degradation: 76ms → 3076ms (40× increase)")

add_heading(doc, "4.5 Scenario D — Resource Exhaustion", 2)
add_para(doc, "POST /api/users/login (Auth — bcrypt bottleneck):")
add_table(doc,
    ["VUs", "OK", "Err%", "p50", "p95", "RPS"],
    [
        ["1",   "1",   "0.0%", "60ms",  "60ms",   "17"],
        ["10",  "10",  "0.0%", "138ms", "197ms",  "51"],
        ["50",  "50",  "0.0%", "478ms", "841ms",  "57"],
        ["100", "100", "0.0%", "899ms", "1636ms", "59"],
    ],
    col_widths=[1.5, 1.5, 2.0, 2.5, 2.5, 2.0]
)
add_para(doc, "GET /api/articles (Feed):")
add_table(doc,
    ["VUs", "OK", "Err%", "p50", "p95", "RPS"],
    [
        ["1",   "1",   "0.0%", "27ms",  "27ms",  "83"],
        ["10",  "10",  "0.0%", "47ms",  "49ms",  "208"],
        ["50",  "50",  "0.0%", "204ms", "209ms", "242"],
        ["100", "100", "0.0%", "330ms", "338ms", "299"],
    ],
    col_widths=[1.5, 1.5, 2.0, 2.5, 2.5, 2.0]
)
add_para(doc, "Result: 0.0% error rate at all concurrency levels ✓  |  Login RPS cap: ~59 (bcrypt CPU-bound)")

# ── 5. Comparative Analysis ───────────────────────────────────────────────────
add_heading(doc, "5. Comparative Analysis: Expected vs Actual", 1)
add_table(doc,
    ["Module", "Midterm Prediction", "Performance Finding", "Mutation Finding", "Chaos Finding"],
    [
        ["Authentication",
         "HIGH — JWT critical path",
         "bcrypt bottleneck confirmed; p95=267ms spike; cap ~59 RPS",
         "100% score slugify; email regex gap",
         "Fails gracefully; ECONNREFUSED <2ms; no retry"],
        ["Article Management",
         "HIGH — core function",
         "Within all thresholds; create p95=69ms spike",
         "100% slugify; appendTagList null-guard gap",
         "Consistent {errors:{}} JSON on DB failure"],
        ["Global Feed",
         "HIGH — main entry",
         "N+1 query: p95=162ms@15VU vs 76ms baseline",
         "paginationParams 85.7%; equivalent mutant",
         "0% errors at 100 VU; graceful latency degradation"],
    ],
    col_widths=[3.0, 3.0, 4.5, 4.0, 4.0]
)
add_para(doc,
    "Key discrepancy: Feed endpoint N+1 query pattern was underestimated in midterm — "
    "only visible under latency injection (40× p95 degradation at 3s injected delay). "
    "Mutation score (90.57%) exceeded baseline expectations.")

# ── 6. Lessons Learned ────────────────────────────────────────────────────────
add_heading(doc, "6. Lessons Learned and Recommendations", 1)

add_heading(doc, "6.1 System Strengths", 2)
for item in [
    "Consistent {errors:{...}} JSON schema on all error responses — no raw stack traces exposed",
    "No request hangs under API downtime — ECONNREFUSED returned in <2ms",
    "0% error rate across all endpoints under 100 concurrent users",
    "Endurance run (900 requests, ~3.3 minutes) — zero latency drift, no memory leak",
    "Mutation score 90.57% exceeds HIGH threshold — test suite detects realistic code faults",
]:
    add_bullet(doc, item)

add_heading(doc, "6.2 Weaknesses and Gaps", 2)
add_table(doc,
    ["Priority", "Area", "Gap", "Impact"],
    [
        ["HIGH",   "Backend",     "No retry logic on DB failure",               "Users get errors on transient DB hiccups"],
        ["HIGH",   "Frontend",    "No error state for 500/503 responses",        "Blank screen or unhandled rejection"],
        ["MEDIUM", "Auth",        "No rate limiting on /api/users/login",        "Brute-force and resource exhaustion possible"],
        ["MEDIUM", "Feed",        "N+1 Sequelize query on article list",         "p95 grows 2.4× from normal to spike load"],
        ["LOW",    "Test suite",  "isValidEmail regex boundaries not tested",    "3 Stryker mutants survive (regex anchors)"],
    ],
    col_widths=[2.0, 2.5, 6.0, 5.5]
)

add_heading(doc, "6.3 Recommendations", 2)
add_table(doc,
    ["#", "Recommendation", "Effort"],
    [
        ["1", "Add express-rate-limit to login: max 10 req/15 min per IP", "Low"],
        ["2", "Add axios interceptor in frontend for 5xx → user-friendly error page", "Low"],
        ["3", "Add 3-attempt retry with exponential backoff for Sequelize operations", "Medium"],
        ["4", "Cache GET /api/tags and unauthenticated GET /api/articles (LRU)", "Medium"],
        ["5", "Add circuit-breaker (opossum) for DB-bound routes", "High"],
        ["6", "Add test cases for leading/trailing whitespace email inputs", "Low"],
    ],
    col_widths=[1.0, 11.5, 2.0]
)

# ── 7. Experimental Setup ─────────────────────────────────────────────────────
add_heading(doc, "7. Experimental Setup", 1)

add_heading(doc, "7.1 Environment", 2)
add_table(doc,
    ["Component", "Details"],
    [
        ["OS",       "Windows 10 (local), Ubuntu Latest (GitHub Actions CI)"],
        ["Node.js",  "20.x"],
        ["Database", "SQLite 3 in-memory (local), PostgreSQL 15 (CI)"],
        ["Backend",  "Express.js 5.2.1 + Sequelize 6.37.8"],
    ],
    col_widths=[4.0, 12.5]
)

add_heading(doc, "7.2 Tools and Versions", 2)
add_table(doc,
    ["Tool", "Version", "Purpose"],
    [
        ["Newman",                        "6.2.2", "Performance test runner"],
        ["@stryker-mutator/core",         "9.6.1", "Mutation testing engine"],
        ["@stryker-mutator/vitest-runner","9.6.1", "Stryker test runner adapter"],
        ["vitest",                        "4.1.5", "Unit test framework for Stryker"],
        ["node:http (built-in)",          "Node.js 20", "Chaos test HTTP client/proxy"],
        ["Playwright",                    "1.59.1", "E2E browser tests (existing)"],
    ],
    col_widths=[5.5, 3.0, 8.0]
)

add_heading(doc, "7.3 Reproduction Commands", 2)
add_para(doc, "Tests that run without a live server:")
for cmd in [
    "npm run test:mutation       # manual mutation runner (10 mutants)",
    "npm run test:stryker        # Stryker (53 mutants, ~8 seconds)",
    "npm run test:chaos:downtime # API downtime simulation",
    "npm run test:chaos:db       # DB failure simulation",
]:
    p = doc.add_paragraph(style="No Spacing")
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(cmd)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK

add_para(doc, "Tests that require backend running:")
for cmd in [
    "npm run test:perf:normal    # 1 VU x 30 iterations",
    "npm run test:perf:peak      # 5 VU x 20 iterations",
    "npm run test:perf:spike     # 15 VU x 5 iterations",
    "npm run test:perf:endurance # 1 VU x 100 iterations",
    "npm run test:chaos:latency  # network latency injection",
    "npm run test:chaos:exhaustion # concurrent flood 1-100 VU",
]:
    p = doc.add_paragraph(style="No Spacing")
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(cmd)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK

doc.add_paragraph()
add_heading(doc, "7.4 CI/CD Integration", 2)
add_table(doc,
    ["Job", "Trigger", "Quality Gate"],
    [
        ["mutation-tests",      "Every push / PR", "QG05: score ≥ 70%"],
        ["chaos-tests",         "Every push / PR", "QG06: all pass"],
        ["performance-tests",   "Manual (workflow_dispatch)", "—"],
    ],
    col_widths=[4.5, 5.5, 6.5]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "docs/export/assignment3-report.docx"
doc.save(out)
print(f"Saved: {out}")
